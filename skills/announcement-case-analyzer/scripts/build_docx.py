#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""上市公司公告案例分析报告 -> Word (.docx) 生成器

用法:
    build_docx.py <input.json> <output.docx>

输入 JSON 结构见 references/report_template.md。
依赖: python-docx (受管虚拟环境已装)
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "Microsoft YaHei"
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)   # 结论强调色（红）
GREEN = RGBColor(0x1E, 0x7E, 0x34)   # 达标绿
WARN = RGBColor(0xB9, 0x6A, 0x00)   # 局限性橙
BOX_BG = "EEF3F9"                     # 速览卡片浅蓝底


def set_run_font(run, size=None, bold=None, color=None, font=CJK):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), font)
    rf.set(qn('w:ascii'), font)
    rf.set(qn('w:hAnsi'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyper = OxmlElement('w:hyperlink')
    hyper.set(qn('r:id'), rid)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyper.append(r)
    paragraph._p.append(hyper)
    for rr in hyper.findall(qn('w:r')):
        rpr = rr.get_or_add_rPr()
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rpr.append(color)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)


def kv_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, bold=True, color=GREY)
    r2 = p.add_run(value or "—")
    set_run_font(r2, size=10.5)
    return p


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F3A5F')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def para(doc, text, size=10.5, bold=False, color=None, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text or "—")
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cell_border(cell, color='1F3A5F', sz='10'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)


def _box_line(cell, icon, label, value, first=False,
              icon_color=NAVY, label_color=NAVY, value_color=None,
              value_bold=False, value_size=10.5):
    """在速览卡片单元格内加一行：图标 + 加粗标签 + 内容。"""
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if icon:
        set_run_font(p.add_run(icon + " "), size=value_size, bold=True, color=icon_color)
    if label:
        set_run_font(p.add_run(label + "："), size=value_size, bold=True, color=label_color)
    if value:
        set_run_font(p.add_run(value), size=value_size, bold=value_bold, color=value_color)
    return p


def summary_box(doc, data, cases):
    """文档最开头的『结论速览』卡片（第一色块）：一眼看懂结论、案例数、关键规律。局限性已移出至独立第三色块。"""
    meta = data.get('meta', {})
    n = len(cases)
    verdict = data.get('verdict') or ''
    if not verdict:
        concl = data.get('conclusion', '') or data.get('overview', '')
        verdict = (concl.split('。')[0] + '。') if '。' in concl else concl[:140]
    findings = (data.get('core_findings') or [])[:3]

    # 标题条
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after = Pt(2)
    set_run_font(tp.add_run("■ 结论速览  EXECUTIVE SUMMARY"), size=12.5, bold=True, color=NAVY)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, BOX_BG)
    set_cell_border(cell, color='1F3A5F', sz='8')

    # 主题
    if meta.get('topic'):
        _box_line(cell, "◆", "主题", meta.get('topic'), first=True)
        first = False
    else:
        first = True

    # 案例数量（带达标/不足标记）
    if n >= 20:
        cnt_txt = "%d 个（满足 ≥20 案例下限）" % n
        cnt_color = GREEN
    else:
        cnt_txt = "%d 个（低于建议下限 20，详见局限性）" % n
        cnt_color = WARN
    _box_line(cell, "●", "案例数", cnt_txt, first=first,
              value_color=cnt_color, value_bold=True)

    # 核心结论（强调）
    _box_line(cell, "★", "核心结论", verdict, value_color=ACCENT,
              value_bold=True, value_size=11)

    # 关键规律（最多3条）
    for f in findings:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.line_spacing = 1.2
        set_run_font(p.add_run("→ "), size=10, bold=True, color=NAVY)
        set_run_font(p.add_run(f), size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def color_block(doc, title, title_en, body_lines, bg_hex,
                title_color=NAVY, body_color=GREY, icon=None):
    """渲染一个带标题的浅底色卡片块（色块）。body_lines 为字符串列表，每条一行。"""
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after = Pt(2)
    set_run_font(tp.add_run("■ %s  %s" % (title, title_en)),
                 size=12.5, bold=True, color=title_color)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    set_cell_border(cell, color='1F3A5F', sz='8')
    first = True
    for ln in body_lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        if icon:
            set_run_font(p.add_run(icon + " "), size=10, bold=True, color=title_color)
        set_run_font(p.add_run(ln), size=10, color=body_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def legal_block(doc, data):
    """第二个色块：法律分析（依据规则解释结果为何如此呈现）。"""
    legal = data.get('legal_analysis') or ''
    if legal:
        color_block(doc, "法律分析", "LEGAL ANALYSIS", [legal],
                    bg_hex="EAF1FB", title_color=NAVY, body_color=GREY)


def limitations_block(doc, data):
    """第三个色块：局限性（如有必填）。"""
    lms = data.get('limitations') or []
    if not lms and data.get('note'):
        lms = [data.get('note')]
    if lms:
        color_block(doc, "局限性", "LIMITATIONS", lms,
                    bg_hex="FDF1E3", title_color=WARN, body_color=GREY, icon="※")


def main():
    if len(sys.argv) < 3:
        print("usage: build_docx.py <input.json> <output.docx>")
        sys.exit(1)
    inp, out = sys.argv[1], sys.argv[2]
    with open(inp, encoding='utf-8') as f:
        data = json.load(f)

    meta = data.get('meta', {})
    cases = data.get('cases', [])
    conclusion = data.get('conclusion', '')
    notes = data.get('notes', '')

    doc = Document()
    nstyle = doc.styles['Normal']
    rpr = nstyle.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), CJK)
    rf.set(qn('w:ascii'), CJK)
    rf.set(qn('w:hAnsi'), CJK)

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t.add_run("上市公司公告案例分析报告"), size=18, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run(meta.get('topic', '')), size=12, bold=True, color=GREY)

    kv_line(doc, "检索范围", meta.get('scope', ''))
    kv_line(doc, "案例类型", meta.get('case_type', ''))
    kv_line(doc, "数据来源", meta.get('mode', ''))
    kv_line(doc, "生成日期", meta.get('generated_at', ''))

    # 速览卡片（第一个色块）
    summary_box(doc, data, cases)
    # 法律分析（第二个色块，置于速览之下）
    legal_block(doc, data)
    # 局限性（第三个色块）
    limitations_block(doc, data)

    # 总览（用户要求：先一段话总览，再给详细分析）
    if data.get('overview'):
        section(doc, "总览")
        _ov = doc.add_paragraph()
        _ov.paragraph_format.space_after = Pt(6)
        _ov.paragraph_format.line_spacing = 1.25
        set_run_font(_ov.add_run(data['overview']), size=10.5)

    # 核心发现（Wind 五步法 Step5：关键信息要点，可选）
    cf = data.get('core_findings') or []
    if cf:
        section(doc, "核心发现")
        for item in cf:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(item), size=10.5)

    # 案例索引表
    section(doc, "一、案例索引")
    if cases:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Light Grid Accent 1'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["#", "公司", "代码", "公告标题", "日期"]):
            hdr[i].text = ''
            set_run_font(hdr[i].paragraphs[0].add_run(h), size=9.5, bold=True, color=NAVY)
        for idx, c in enumerate(cases, 1):
            row = tbl.add_row().cells
            vals = [str(idx), c.get('company', ''), c.get('code', ''), c.get('title', ''), c.get('date', '')]
            for i, v in enumerate(vals):
                row[i].text = ''
                set_run_font(row[i].paragraphs[0].add_run(v), size=9)

    # 逐案例
    for idx, c in enumerate(cases, 1):
        section(doc, "案例%d　%s（%s）— %s" % (idx, c.get('company', ''), c.get('code', ''), c.get('title', '')))
        kv_line(doc, "基本要素", "交易所：%s　|　日期：%s　|　类型：%s" % (
            c.get('exchange', ''), c.get('date', ''), c.get('type', '')))
        para(doc, "【案例概要】", size=10.5, bold=True, color=NAVY, space=2)
        para(doc, c.get('summary', ''), size=10.5, space=6)
        para(doc, "【案例分析】", size=10.5, bold=True, color=NAVY, space=2)
        para(doc, c.get('analysis', ''), size=10.5, space=6)
        if c.get('regulatory'):
            para(doc, "【监管视角 / 可借鉴点】", size=10.5, bold=True, color=NAVY, space=2)
            para(doc, c.get('regulatory', ''), size=10.5, space=6)
        quotes = c.get('quotes') or []
        if quotes:
            para(doc, "【引用片段】", size=10.5, bold=True, color=NAVY, space=2)
            for q in quotes:
                qp = doc.add_paragraph()
                qp.paragraph_format.left_indent = Pt(18)
                qp.paragraph_format.space_after = Pt(3)
                qp.paragraph_format.line_spacing = 1.15
                r = qp.add_run("“%s”" % q)
                set_run_font(r, size=9.5, color=GREY)
        if c.get('link'):
            p = doc.add_paragraph()
            set_run_font(p.add_run("原文链接："), size=10.5, bold=True, color=GREY)
            add_hyperlink(p, c['link'], c['link'])

    if conclusion:
        section(doc, "二、总体结论")
        para(doc, conclusion, size=10.5, space=6)
    if notes:
        section(doc, "附注")
        para(doc, notes, size=9.5, color=GREY, space=4)

    doc.save(out)
    print("OK", out, "cases=", len(cases))


if __name__ == '__main__':
    main()
